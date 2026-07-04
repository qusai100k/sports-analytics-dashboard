from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "documentation" / "report_figures"
TABLE_DIR = ROOT / "documentation" / "report_tables"
OUT = ROOT / "Final_Thesis.docx"
PDF_OUT = ROOT / "Final_Thesis.pdf"


with (TABLE_DIR / "report_facts.json").open(encoding="utf-8") as f:
    FACTS = json.load(f)


FIGURES = [
    ("Figure 4.1", "Home page of the Sports Analytics Dashboard.", "figure_4_1_home.png"),
    ("Figure 4.2", "Dataset Overview page showing dataset statistics and table preview.", "figure_4_2_dataset_overview.png"),
    ("Figure 4.3", "Correlation heatmap generated from the football dataset.", "figure_4_3_correlation_heatmap.png"),
    ("Figure 4.4", "Data Visualization page with interactive charts.", "figure_4_4_data_visualization.png"),
    ("Figure 4.5", "Team Comparison page with head-to-head and performance charts.", "figure_4_5_team_comparison.png"),
    ("Figure 4.6", "Match Prediction page with league and team selection.", "figure_4_6_match_prediction_form.png"),
    ("Figure 4.7", "Prediction output with probability and explanation.", "figure_4_7_prediction_output.png"),
    ("Figure 4.8", "World Cup Prediction page with team strength index.", "figure_4_8_world_cup_prediction.png"),
    ("Figure 4.9", "Model Performance page showing evaluation metrics.", "figure_4_9_model_performance.png"),
    ("Figure 4.10", "Confusion matrix section from the Model Performance page.", "figure_4_10_confusion_matrix.png"),
    ("Figure 4.11", "Feature importance section from the Model Performance page.", "figure_4_11_feature_importance.png"),
    ("Figure 4.12", "About Project page explaining the technology stack and pipeline.", "figure_4_12_about_project.png"),
]

TABLES = [
    ("Table 3.1", "Dataset summary used in the project."),
    ("Table 3.2", "Main engineered features used by the prediction pipeline."),
    ("Table 4.1", "Model evaluation metrics."),
    ("Table 4.2", "Top feature importance values."),
    ("Table 4.3", "Confusion matrix from the saved model bundle."),
    ("Table C.1", "Sample rows from the football dataset."),
]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_body(doc: Document, text: str):
    p = add_paragraph(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_caption(doc: Document, label: str, text: str):
    p = add_paragraph(doc, f"{label}. {text}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(10)


def add_figure(doc: Document, label: str, caption: str, filename: str):
    path = FIG_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.95))
    add_caption(doc, label, caption)


def add_table(doc: Document, label: str, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    add_caption(doc, label, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(11, 37, 69)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    return doc


def cover_page(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    title = add_paragraph(doc, "Sports Analytics Dashboard and Match Outcome Prediction Using Machine Learning", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    for run in title.runs:
        run.font.size = Pt(20)
        run.bold = True
    doc.add_paragraph()
    add_paragraph(doc, "Mini Skripsi", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Undergraduate Final Project Report", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "Prepared for University Final Project Assessment", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Sports Analytics and Machine Learning", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def front_matter(doc: Document):
    doc.add_heading("Approval Page", level=1)
    add_body(doc, "This mini thesis entitled Sports Analytics Dashboard and Match Outcome Prediction Using Machine Learning has been prepared as an undergraduate final project report. The work describes the implementation, evaluation, and results of the sports analytics platform developed in the repository.")
    add_body(doc, "This report is submitted for academic review and final project assessment.")
    doc.add_paragraph("\n\n")
    add_body(doc, "Approved by:")
    doc.add_paragraph("\n")
    add_body(doc, "Academic Supervisor: ______________________________")
    add_body(doc, "Date: ______________________________")
    doc.add_page_break()

    doc.add_heading("Declaration", level=1)
    add_body(doc, "I declare that this mini thesis is based on the implementation found in the Sports Analytics Dashboard project repository. The dataset description, model results, screenshots, and discussion are taken from the actual application files, generated outputs, and saved model bundle. No experimental result or screenshot has been invented.")
    add_body(doc, "All external references are listed in APA style in the References section.")
    doc.add_page_break()

    doc.add_heading("Acknowledgements", level=1)
    add_body(doc, "I would like to thank the academic supervisor and lecturers who provided guidance during this final project. I also appreciate the open-source communities behind Python, pandas, scikit-learn, Streamlit, Plotly, and OpenFootball, because these tools and datasets made the project possible.")
    add_body(doc, "This project helped me understand how software engineering, data analysis, machine learning, and user interface design can be combined into one complete application.")
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_body(doc, "Football is one of the most popular sports in the world, and match data can be used to understand team performance and support prediction. This project developed a Sports Analytics Dashboard and Match Outcome Prediction system using Python, Streamlit, Plotly, pandas, NumPy, scikit-learn, Joblib, and Pillow. The application uses verified football data from the OpenFootball public dataset cache and stores a processed dataset containing 684 match records, 27 columns, and 23 real football clubs.")
    add_body(doc, f"The system includes dataset overview, data visualization, team comparison, match prediction, World Cup prediction, model performance, and project information pages. The saved model bundle uses {FACTS['model_name']} for match outcome classification. The model reached an accuracy of {metric('Accuracy')}%, precision of {metric('Precision')}%, recall of {metric('Recall')}%, and F1 score of {metric('F1')}% on the test split. The prediction page also supports verified teams, optional live API providers, cached data, probability output, confidence score, and explanation text.")
    add_body(doc, "The result is a working undergraduate capstone application that demonstrates the full process of collecting football data, preparing features, training a model, evaluating performance, and presenting the results through a modern web dashboard.")
    add_body(doc, "Keywords: sports analytics, football analytics, machine learning, Streamlit, match prediction, data visualization")
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Approval Page", "Declaration", "Acknowledgements", "Abstract", "List of Figures", "List of Tables",
        "Chapter 1 Introduction", "Chapter 2 Literature Review", "Chapter 3 System Design and Methodology",
        "Chapter 4 Implementation and Results", "Chapter 5 Conclusion", "References", "Appendices",
    ]
    for item in toc_items:
        add_body(doc, item)
    doc.add_page_break()

    doc.add_heading("List of Figures", level=1)
    for label, caption, _ in FIGURES:
        add_body(doc, f"{label}. {caption}")
    doc.add_page_break()

    doc.add_heading("List of Tables", level=1)
    for label, caption in TABLES:
        add_body(doc, f"{label}. {caption}")
    doc.add_page_break()


def metric(name: str) -> str:
    for row in FACTS["metrics"]:
        if row["Metric"] == name:
            return f"{float(row['Value (%)']):.2f}"
    return "0.00"


def chapter1(doc: Document):
    doc.add_heading("CHAPTER 1 INTRODUCTION", level=1)
    doc.add_heading("1.1 Background", level=2)
    add_body(doc, "Football produces a large amount of data from every match, such as goals, results, team form, and performance trends. In the past, this information was often read manually from tables or match reports. Today, sports analytics allows the same information to be processed and displayed in a more useful way. A dashboard can help users see patterns faster, and machine learning can be used to estimate possible match outcomes.")
    add_body(doc, "This project focuses on building a complete football analytics platform. It is not only a model script, and it is not only a dashboard. The system combines dataset processing, data visualization, team comparison, prediction, model evaluation, and a World Cup strength index into one Streamlit web application.")
    doc.add_heading("1.2 Problem Statement", level=2)
    add_body(doc, "Football data is available from different sources, but it is not always easy for students or general users to explore it. Raw datasets can be difficult to read, and prediction outputs can be hard to understand if they only show a class label without probability or explanation. Another problem is that many simple dashboard projects use artificial data, which makes the application less realistic.")
    add_body(doc, "The problem addressed in this project is how to build a realistic, usable, and explainable sports analytics dashboard using verified football data and a machine learning pipeline.")
    doc.add_heading("1.3 Research Questions", level=2)
    add_bullets(doc, [
        "How can verified football match data be organized into an interactive analytics dashboard?",
        "Which football features can be used to support match outcome prediction?",
        "How well does the saved machine learning model perform on the available dataset?",
        "How can prediction results be explained in simple English for dashboard users?",
    ])
    doc.add_heading("1.4 Objectives", level=2)
    add_bullets(doc, [
        "To develop a Streamlit dashboard for football dataset exploration and visualization.",
        "To create a team comparison feature using real football clubs and performance indicators.",
        "To build a match prediction workflow that returns class probabilities and confidence score.",
        "To evaluate the model using accuracy, precision, recall, F1 score, confusion matrix, and feature importance.",
        "To include a World Cup prediction page based on a transparent weighted strength index.",
    ])
    doc.add_heading("1.5 Scope", level=2)
    add_body(doc, "The project uses the existing repository as the source of truth. The main dataset is the processed football match dataset saved in dataset/football_matches.csv. The dashboard supports the implemented pages and the optional live provider structure already present in the code. The model result discussion is limited to the saved model bundle and does not claim performance beyond the available dataset.")
    doc.add_heading("1.6 Benefits", level=2)
    add_body(doc, "The project benefits students by showing how a machine learning model can be connected to a real web application. It also helps users understand football data through charts, tables, and simple prediction explanations. From a software engineering point of view, it demonstrates caching, modular code, reusable components, and separation between data utilities, model logic, visualization, and pages.")
    doc.add_heading("1.7 Report Organization", level=2)
    add_body(doc, "Chapter 1 introduces the background, problem, questions, objectives, scope, and benefits. Chapter 2 reviews sports analytics, football analytics, machine learning, dashboards, and related tools. Chapter 3 explains the system design and methodology. Chapter 4 presents the implementation and results from the actual application. Chapter 5 concludes the report and explains limitations and future work.")
    doc.add_page_break()


def chapter2(doc: Document):
    doc.add_heading("CHAPTER 2 LITERATURE REVIEW", level=1)
    sections = [
        ("2.1 Sports Analytics", "Sports analytics is the use of data to understand performance, strategy, and decision-making in sports. In football, analytics can support scouting, team preparation, post-match review, and prediction. The value of analytics is not only in collecting data but also in turning data into useful insight."),
        ("2.2 Football Analytics", "Football analytics focuses on match events, team performance, scoring patterns, defensive strength, and player or team behavior. Public football datasets make it possible for students to build realistic projects without depending only on private professional data."),
        ("2.3 Machine Learning", "Machine learning allows a computer program to learn patterns from historical data. In this project, the task is classification because the model predicts one of three possible outcomes: Home Win, Draw, or Away Win."),
        ("2.4 Data Visualization", "Data visualization helps users understand complex data quickly. Charts such as histograms, boxplots, heatmaps, radar charts, and bar charts are useful because they show patterns that are difficult to see in raw tables."),
        ("2.5 Random Forest", "Random Forest is an ensemble method introduced by Breiman (2001). It combines many decision trees and is widely used for classification tasks. In this project, Random Forest is included in the model comparison pipeline, while the saved best model in the current bundle is Logistic Regression."),
        ("2.6 Streamlit Dashboard", "Streamlit is a Python framework for building data applications. It is useful for undergraduate projects because it allows the developer to connect Python analysis, charts, machine learning models, and user interface components in one application."),
        ("2.7 Match Prediction Systems", "A match prediction system estimates the possible result of a future or selected match. A good system should not only show the predicted outcome but should also show probability, confidence, and explanation so the user understands the reason behind the result."),
        ("2.8 Related Studies", "Previous work in sports prediction shows that machine learning can support result prediction when meaningful features are available. However, model performance depends strongly on dataset quality, feature design, and the nature of the sport. Football is difficult to predict because draws, injuries, tactics, and random events can influence the final result."),
    ]
    for title, text in sections:
        doc.add_heading(title, level=2)
        add_body(doc, text)
    add_body(doc, "The literature supports the design direction of this project: use verified data, prepare meaningful features, evaluate the model honestly, and present the results through a dashboard that is easy to understand.")
    doc.add_page_break()


def chapter3(doc: Document):
    doc.add_heading("CHAPTER 3 SYSTEM DESIGN AND METHODOLOGY", level=1)
    doc.add_heading("3.1 Project Overview", level=2)
    add_body(doc, "The project is a Python-based web application built with Streamlit. The application is organized into app.py, pages, components, assets, models, dataset, and utils folders. The pages folder contains dashboard pages, while utils contains data loading, live provider handling, model utilities, report generation, visualization utilities, training logic, and World Cup indexing.")
    doc.add_heading("3.2 Dataset", level=2)
    rows = [[str(r["Item"]), str(r["Value"])] for r in FACTS["dataset_summary"]]
    add_table(doc, "Table 3.1", "Dataset summary used in the project.", ["Item", "Value"], rows, [2.2, 4.1])
    add_body(doc, "The dataset summary in Table 3.1 comes directly from dataset/football_matches.csv. The project currently contains real Premier League data from the OpenFootball public cache.")
    doc.add_heading("3.3 Data Collection", level=2)
    add_body(doc, "The data collection logic is implemented in the live data utilities. The priority is to use optional live APIs when API keys are available, then cached API data, then OpenFootball public data, then football-data.co.uk public data, and finally verified real team lists. The application does not fabricate teams or fixtures.")
    doc.add_heading("3.4 Data Cleaning", level=2)
    add_body(doc, "The project parses football match data, standardizes team names, removes duplicate match rows, converts dates, and ensures that required columns such as home team, away team, goals, result, season, and league are available. Missing values and duplicate rows are shown on the Dataset Overview page.")
    doc.add_heading("3.5 Feature Engineering", level=2)
    feature_rows = [
        ["home_team_strength", "Estimated home team strength before the match."],
        ["away_team_strength", "Estimated away team strength before the match."],
        ["home_recent_form", "Recent form signal for the home team."],
        ["away_recent_form", "Recent form signal for the away team."],
        ["home_avg_goals", "Average goals signal for the home team."],
        ["away_avg_goals", "Average goals signal for the away team."],
        ["home_defense_rating", "Defensive rating for the home team."],
        ["away_defense_rating", "Defensive rating for the away team."],
        ["team_strength_delta", "Difference between home and away team strength."],
        ["goal_power_delta", "Difference between attacking goal signals."],
    ]
    add_table(doc, "Table 3.2", "Main engineered features used by the prediction pipeline.", ["Feature", "Description"], feature_rows, [2.5, 3.8])
    doc.add_heading("3.6 Machine Learning Pipeline", level=2)
    add_body(doc, "The machine learning pipeline is implemented in utils/train_model.py. It loads verified football match data, engineers model features, encodes the target variable, splits the dataset into training and testing sets, trains candidate models, compares metrics, and saves the selected bundle using Joblib.")
    doc.add_heading("3.7 Model Training", level=2)
    add_body(doc, f"The saved model bundle reports {FACTS['model_name']} as the selected model. The target variable is result, with the classes Away Win, Draw, and Home Win. The training split contains {FACTS['training_info']['training_rows']} rows and the testing split contains {FACTS['training_info']['testing_rows']} rows.")
    doc.add_heading("3.8 Model Evaluation", level=2)
    add_body(doc, "The model is evaluated using accuracy, precision, recall, F1 score, cross-validation mean and standard deviation, classification report, confusion matrix, ROC curve, and feature importance. These values are displayed on the Model Performance page.")
    doc.add_heading("3.9 Dashboard Development", level=2)
    add_body(doc, "The dashboard uses Streamlit for the web interface and Plotly for interactive charts. Shared styling and navigation are handled in components/ui.py. This keeps the visual identity consistent across pages.")
    doc.add_heading("3.10 Prediction Workflow", level=2)
    add_body(doc, "For club matches, the user selects a league and either chooses an available upcoming fixture or manually selects two verified real teams. The system calculates team statistics, recent form, team strength, and probabilities. For World Cup prediction, national teams are compared using a weighted strength index.")
    doc.add_heading("3.11 System Architecture", level=2)
    add_body(doc, "The system workflow is: verified data source, data cleaning, feature engineering, model training, prediction utility, Streamlit dashboard, and downloadable reports. This architecture separates data processing, model logic, visual components, and page-level user interface code.")
    doc.add_page_break()


def chapter4(doc: Document):
    doc.add_heading("CHAPTER 4 IMPLEMENTATION AND RESULTS", level=1)
    doc.add_heading("4.1 Implementation Overview", level=2)
    add_body(doc, "The implementation is divided into reusable Python modules. The app.py file provides the home page. The pages folder contains the user-facing dashboard pages. The utils folder contains data loading, live provider logic, model prediction, report generation, visualization functions, model training, World Cup calculations, and presentation content helpers.")
    add_figure(doc, *FIGURES[0])
    doc.add_heading("4.2 Dataset Overview Page", level=2)
    add_body(doc, "Figure 4.2 shows the Dataset Overview page developed in this project. It displays dataset size, missing values, duplicate values, an interactive data table, dataset summary, feature types, and a download button.")
    add_figure(doc, *FIGURES[1])
    add_body(doc, "Figure 4.3 shows the correlation heatmap. This chart helps identify how numerical features relate to each other before model training and evaluation.")
    add_figure(doc, *FIGURES[2])
    doc.add_heading("4.3 Data Visualization Page", level=2)
    add_body(doc, "Figure 4.4 shows the Data Visualization page. The page includes interactive charts such as outcome distribution, correlation matrix, histograms, boxplots, scatter plots, and violin charts. Users can filter by season, league, and team.")
    add_figure(doc, *FIGURES[3])
    doc.add_heading("4.4 Team Comparison", level=2)
    add_body(doc, "Figure 4.5 shows the Team Comparison page. It allows the user to choose two real football clubs and compare win rate, goals, average performance, radar chart, bar charts, head-to-head data, and performance trend.")
    add_figure(doc, *FIGURES[4])
    doc.add_heading("4.5 Match Prediction", level=2)
    add_body(doc, "Figure 4.6 shows the Match Prediction page. The page includes league selection, upcoming match selection when verified fixtures are available, and manual team selection using only real teams.")
    add_figure(doc, *FIGURES[5])
    add_body(doc, "Figure 4.7 shows a prediction output generated from the running application. The output contains the predicted outcome, confidence score, provider badge, recent form cards, strength comparison, probability bars, and explanation.")
    add_figure(doc, *FIGURES[6])
    doc.add_heading("4.6 World Cup Prediction", level=2)
    add_body(doc, "Figure 4.8 shows the World Cup Prediction page. It uses a transparent weighted formula: official senior win is 1.00, draw is 0.50, friendly or non-official win is 0.30, and youth signal is 0.20. The result is shown as a team strength index.")
    add_figure(doc, *FIGURES[7])
    doc.add_heading("4.7 Model Performance", level=2)
    metric_rows = [[r["Metric"], f"{float(r['Value (%)']):.2f}"] for r in FACTS["metrics"]]
    add_table(doc, "Table 4.1", "Model evaluation metrics.", ["Metric", "Value (%)"], metric_rows, [3.0, 3.0])
    add_body(doc, f"The saved model is {FACTS['model_name']}. Table 4.1 reports the actual metrics from the saved model bundle. The model reached {metric('Accuracy')}% accuracy and {metric('F1')}% F1 score. These values show that the model works as a baseline classifier but still has room for improvement.")
    add_figure(doc, *FIGURES[8])
    add_body(doc, "Figure 4.10 shows the confusion matrix section from the Model Performance page. The confusion matrix helps show which outcome classes are easier or harder for the model to classify.")
    add_figure(doc, *FIGURES[9])
    importance_rows = [[r["feature"], str(r["importance"])] for r in FACTS["feature_importance"][:8]]
    add_table(doc, "Table 4.2", "Top feature importance values.", ["Feature", "Importance"], importance_rows, [3.2, 2.3])
    add_body(doc, "Table 4.2 and Figure 4.11 show the most important features from the current saved model bundle. The highest ranked feature is away_defense_rating, followed by home_team_strength and home_defense_rating.")
    add_figure(doc, *FIGURES[10])
    cm_headers = ["Actual"] + FACTS["labels"]
    cm_rows = []
    for row in FACTS["confusion_matrix"]:
        cm_rows.append([row["Actual"]] + [str(row[label]) for label in FACTS["labels"]])
    add_table(doc, "Table 4.3", "Confusion matrix from the saved model bundle.", cm_headers, cm_rows, [1.3, 1.5, 1.5, 1.5])
    doc.add_heading("4.8 About Page and Project Explanation", level=2)
    add_body(doc, "Figure 4.12 shows the About Project page. This page explains the project description, technology stack, machine learning pipeline, data source, challenges, future improvements, and university context.")
    add_figure(doc, *FIGURES[11])
    doc.add_heading("4.9 Discussion", level=2)
    add_body(doc, "The implementation shows that a complete sports analytics system can be built using Python and Streamlit. The main strength of the system is that it connects real football data, interactive analysis, and prediction output in one application. The dashboard also avoids fake teams and fabricated fixtures by using verified sources, caches, and real fallback lists.")
    add_body(doc, "The weakness is that the current model performance is moderate. Football is difficult to predict because results can be affected by tactical changes, player injuries, red cards, weather, and other factors that are not fully represented in the current dataset. The model should therefore be treated as an educational baseline rather than a professional betting system.")
    doc.add_page_break()


def chapter5(doc: Document):
    doc.add_heading("CHAPTER 5 CONCLUSION", level=1)
    doc.add_heading("5.1 Conclusion", level=2)
    add_body(doc, "This project successfully developed a Sports Analytics Dashboard and Match Outcome Prediction system using machine learning. The application includes dataset overview, visualization, team comparison, match prediction, World Cup prediction, model performance, and project explanation pages. It uses real football data from the OpenFootball public cache and stores a trained model bundle for prediction and evaluation.")
    add_body(doc, f"The saved model achieved {metric('Accuracy')}% accuracy and {metric('F1')}% F1 score on the test split. The result is suitable for an undergraduate machine learning dashboard project because it shows the complete process from data collection to deployment, while still reporting the limitations honestly.")
    doc.add_heading("5.2 Limitations", level=2)
    add_bullets(doc, [
        "The dataset is limited to the currently cached public match records.",
        "The model does not include injuries, lineups, tactical formations, weather, or betting market signals.",
        "Upcoming fixtures depend on verified providers and are not fabricated when unavailable.",
        "The World Cup prediction uses an indexed strength formula rather than a full tournament simulation.",
    ])
    doc.add_heading("5.3 Future Work", level=2)
    add_bullets(doc, [
        "Connect live API standings, fixtures, and form when API keys are available.",
        "Add player-level statistics, injuries, transfers, and expected goals.",
        "Use stronger time-series validation and more seasons of real match data.",
        "Deploy the application online with authentication and persistent prediction history.",
        "Build a mobile-friendly version for easier use during presentations or match analysis.",
    ])
    doc.add_page_break()


def references(doc: Document):
    doc.add_heading("REFERENCES", level=1)
    refs = [
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32. https://doi.org/10.1023/A:1010933404324",
        "Bunker, R. P., & Thabtah, F. (2019). A machine learning framework for sport result prediction. Applied Computing and Informatics, 15(1), 27-33. https://doi.org/10.1016/j.aci.2017.09.005",
        "OpenFootball Project. (2026). England football data. https://github.com/openfootball/england",
        "Pappalardo, L., Cintia, P., Rossi, A., Massucco, E., Ferragina, P., Pedreschi, D., & Giannotti, F. (2019). A public data set of spatio-temporal match events in soccer competitions. Scientific Data, 6, 236. https://doi.org/10.1038/s41597-019-0247-7",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Plotly Technologies Inc. (2026). Plotly Python graphing library. https://plotly.com/python/",
        "Python Software Foundation. (2026). Python programming language. https://www.python.org/",
        "Snowflake Inc. (2026). Streamlit documentation. https://docs.streamlit.io/",
        "The pandas development team. (2026). pandas documentation. https://pandas.pydata.org/docs/",
        "football-data.org. (2026). Football data API. https://www.football-data.org/",
    ]
    for ref in refs:
        p = add_body(doc, ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    doc.add_page_break()


def appendices(doc: Document):
    doc.add_heading("APPENDICES", level=1)
    doc.add_heading("Appendix A Website Screenshots", level=2)
    add_body(doc, "Appendix A contains selected screenshots captured from the actual running Streamlit application. The complete figure set is stored in documentation/report_figures/.")
    for label, caption, filename in FIGURES[:4]:
        add_figure(doc, label.replace("Figure", "Appendix Figure"), caption, filename)
    doc.add_page_break()

    doc.add_heading("Appendix B Important Code Snippets", level=2)
    snippets = [
        ("Data loading", "utils/data.py", "load_dataset() reads dataset/football_matches.csv and parses the date column."),
        ("Prediction provider", "utils/live_football_data.py", "get_real_teams(), get_upcoming_matches(), get_team_statistics(), and predict_match() provide the real-football prediction flow."),
        ("Model training", "utils/train_model.py", "The training script loads verified public data, engineers features, trains models, and saves the Joblib bundle."),
        ("World Cup index", "utils/world_cup.py", "The World Cup page uses weighted signals for official wins, draws, friendly wins, and youth signals."),
    ]
    add_table(doc, "Table B.1", "Important code snippets and their roles.", ["Area", "File", "Role"], snippets, [1.5, 2.0, 2.7])
    doc.add_page_break()

    doc.add_heading("Appendix C Dataset Description", level=2)
    sample = []
    for row in pd_read(TABLE_DIR / "dataset_sample.csv"):
        sample.append([row["date"], row["home_team"], row["away_team"], row["result"]])
    add_table(doc, "Table C.1", "Sample rows from the football dataset.", ["Date", "Home Team", "Away Team", "Result"], sample, [1.3, 1.8, 1.8, 1.2])
    add_body(doc, "The dataset includes match date, season, league, home team, away team, goals, result, engineered strength features, recent form, and data source.")
    doc.add_page_break()

    doc.add_heading("Appendix D Prediction Examples", level=2)
    add_body(doc, "Prediction examples are generated inside the Match Prediction page by selecting a league, choosing two verified teams, and pressing Predict Match Outcome. The application stores the prediction history in Streamlit session state and allows the result to be exported as a PDF report.")
    add_figure(doc, "Appendix Figure D.1", "Prediction output example captured from the running application.", "figure_4_7_prediction_output.png")


def pd_read(path: Path):
    import csv
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def main() -> None:
    doc = setup_doc()
    cover_page(doc)
    front_matter(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    references(doc)
    appendices(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
